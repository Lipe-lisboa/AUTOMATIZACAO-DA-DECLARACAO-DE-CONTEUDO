import streamlit as st
from docx import Document
from os import path
import os
from datetime import datetime
import sqlite3
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import pandas as pd
import pickle
from subprocess import Popen

# Nome do arquivo para salvar os dados localmente
SESSION_FILE = "session_data.pkl"

# Função para salvar o estado
def save_session():
    # Salva os estados de "dados" e "dados_tabela" no arquivo
    session_data = {
        "dados": st.session_state.get("dados", []),
        "dados_tabela": st.session_state.get("dados_tabela", {
            "Produto": [],
            "Modelo": [],
            "Quantidade": [],
            "Valor": []
        }),
        "peso": st.session_state.get("peso", 1),
    }
    with open(SESSION_FILE, "wb") as f:
        pickle.dump(session_data, f)

# Função para carregar o estado
def load_session():
    
    #verifica se o arquivo existe (quando o usuari gera uma nova dec o arquivo deixa de existir)
    if os.path.exists(SESSION_FILE):
        if os.path.getsize(SESSION_FILE) > 0:  # Verifica se o arquivo não está vazio
            with open(SESSION_FILE, "rb") as f:
                try:
                    return pickle.load(f)
                except EOFError:
                    return {
                        "dados": [], 
                        "dados_tabela": {
                            "Produto": [],
                            "Modelo": [],
                            "Quantidade": [],
                            "Valor": [],
                        },
                        "peso": 1
                    }
    return {
        "dados": [],
        "dados_tabela": {
            "Produto": [],
            "Modelo": [],
            "Quantidade": [],
            "Valor": []
        },
        "peso": 1
    }

# Inicializar a sessão carregando os dados salvos
if "dados" not in st.session_state or "dados_tabela" not in st.session_state or "peso" not in  st.session_state:
    session_data = load_session()
    st.session_state["dados"] = session_data["dados"]
    st.session_state["dados_tabela"] = session_data["dados_tabela"]
    st.session_state["peso"] = session_data["peso"]


# Cria ou abre um banco de dados chamado 'meu_banco.db'
DB_PATH = os.getenv("DB_PATH", "logistica.db")  # Configuração para deploy
conn = sqlite3.connect(DB_PATH)

# Cria um cursor para executar comandos SQL
cursor = conn.cursor()

# Consultar os dados de nome e CNPJ da tabela 'remetentes'
cursor.execute("SELECT nome FROM remetente")

# Recuperar todos os resultados da consulta
remetentes = cursor.fetchall()

# Consultar os dados de nome e CNPJ da tabela 'remetentes'
cursor.execute("SELECT nome FROM destinatario")

# Recuperar todos os resultados da consulta
destinatarios = cursor.fetchall()

#TITULO DA PAGINA
st.title('📦Declaração de Conteudo')

#OUTRAS SESSIONS STATE
if "gerar_dec_conteudo" not in st.session_state:
    st.session_state["gerar_dec_conteudo"] = False
    
if "nova_dec" not in st.session_state:
    st.session_state["nova_dec"] = False
    
if "removido" not in st.session_state:
    st.session_state["removido"] = False

# FUNÇÃO PARA EXCLUIR DADOS (RESET)
def excluir_dados():
    if os.path.exists(SESSION_FILE):
        os.remove(SESSION_FILE)
            
    data_atual = datetime.now()
    dia = data_atual.day
    mes = data_atual.month
    ano = data_atual.year
    
    caminho_dec_docx = f"DECS/DEC_CONTEÚDO_CORREIOS-{dia}_{mes}_{ano}.docx"
    caminho_dec_pdf = caminho_dec_docx.replace('.docx', '.pdf')
    
    
    if os.path.exists(caminho_dec_docx) and os.path.exists(caminho_dec_pdf) :
        os.remove(caminho_dec_docx)
        os.remove(caminho_dec_pdf)
    
    st.session_state["dados"] = []
    st.session_state["gerar_dec_conteudo"] = False
    st.session_state["peso"] = 1
    st.session_state["nova_dec"] = False
    for i in st.session_state["dados_tabela"]:
        st.session_state["dados_tabela"][i] = []
        
    st.rerun()

# FUNÇÃO PARA REMOVER PRODUTO DA DEC
def remover_produto(indice):
    
    try: 
        st.session_state["dados"].pop(indice)
        
        # Remover produto de "dados_tabela"
        st.session_state["dados_tabela"]["Produto"].pop(indice)
        st.session_state["dados_tabela"]["Modelo"].pop(indice)
        st.session_state["dados_tabela"]["Quantidade"].pop(indice)
        st.session_state["dados_tabela"]["Valor"].pop(indice)
            
        save_session()  # Salvar o estado atualizado
            
        st.session_state["removido"] = True
        st.rerun()
    except IndexError:
        st.error('Esse produto não existe, por favor insira um indice existente❗')


# NOVA DEC == FALSE:
if st.session_state["nova_dec"] == False:
    
    #Formulario:
    with st.form('forme_dec'):
        st.title('📋 Preenchimento')
        
        #Pega os remetentes e os Destinatarios no banco de dados:
        lista_remetentes = []
        for remetente in remetentes:
            lista_remetentes.append(f'{remetente[0]}')
            
        lista_destinatarios = []
        for destinatario in destinatarios:
            lista_destinatarios.append(f'{destinatario[0]}')
        
        #coloca os rem e dest no formulario para o usuario poder escolher:
        remetente_selecionado =  st.selectbox('Remetente:', lista_remetentes)
        destinatario_selecionado = st.selectbox('Destinatario:', lista_destinatarios)
        
        #Parte do fomes de add os produtos 
        st.subheader('Adicionar Produto:')
        tipo = st.text_input('Tipo de produto:', placeholder='Exemplo: Caixa de som')
        modelo = st.text_input('Modelo:', placeholder='Exemplo: IMS-123')
        produto_sem_modelo = st.checkbox('Produto não possui modelo')
        qtd = st.number_input('Quantidade:',min_value=1)
        valor = st.number_input('Valor:', min_value=50, step=50)
        
        btt_add_produto = st.form_submit_button('Adicionar Produto')
        
        # VERIFICAÇÃO DO FORMULARIO
        def verificacao(produto):
            result = True
            for i in [tipo, qtd, valor]:
                if not i:
                    print(i)
                    st.error(f'Informações insuficientes. Por favor preencha todas as etapas do formulario❗')
                    result = False
                    
            if not modelo and not produto_sem_modelo:
                st.warning(f'Nenhum modelo foi informado. Caso o produto não tenha modelo\
                         marque a opção "Produto não possui modelo"⚠️')
                result = False
                
            if modelo and produto_sem_modelo:
                st.warning(f'Foi informado um modelo ao produto, porem a opção \
                "Produto não possui modelo" esta marcada.⚠️')
                result = False
                
            if produto in st.session_state["dados"]:
                st.warning('Esse produto ja foi inserido na DEC⚠️')
                result = False
                
            return result

        # Se o botão de add produto for apertado:
        if btt_add_produto:
            if produto_sem_modelo:
                conteudo = f'{tipo}'
            else:
                conteudo = f'{tipo} - Modelo: {modelo}'
                
            produto = {
                "conteudo": conteudo,
                "quantidade": qtd,
                "valor": valor,
            }
            
            verificado = verificacao(produto)
            
            #SE NÃO TIVER NENHUM ERRO NO FORMULARIO:
            if verificado:
                #ADICIONA OS DADOS NA TABELA
                st.session_state["dados_tabela"]['Produto'].append(tipo)
                st.session_state["dados_tabela"]['Modelo'].append(modelo)
                st.session_state["dados_tabela"]['Quantidade'].append(qtd)
                st.session_state["dados_tabela"]['Valor'].append(f'{valor:.2f}R$')
                
                #ADICIONA OS DADOS NA SESSION STATE
                st.session_state["dados"].append(produto)
                save_session()  # Salvar o estado atualizado
                st.success("Produto adicionado com sucesso!!✅")
                
                # ESCREVE AS INFORMAÇÕES ENVIADAS
                st.write(produto['conteudo'])
                st.write(f'Quantidade: {produto["quantidade"]}')
                st.write(f'Valor: {produto["valor"]}')
                
                if st.form_submit_button('Adicionar outro Produto'):
                    st.rerun()
        
        # PARTE DO FORMULARIO QUE ADICIONA O PESO DA CAIXA
        st.subheader('Adicionar Peso:')
        peso = st.number_input('Peso da Caixa (Kg)', min_value=1)
        btt_peso = st.form_submit_button('Adicionar Peso da Caixa')
        
        if btt_peso:
            st.session_state["peso"] = peso
            save_session()  # Salvar o estado atualizado
            st.success('Peso adicionado com sucesso!!✅')
    
    # GERAR/EXCLUIR DEC DE CONTEUDO
    if st.session_state["dados"]:
        st.subheader('Gerar/Excluir DEC')
        
        #GERAR DEC DE CONTEUDO
        btt_gerar_doc = st.button('Gerar Declaração de Conteudo') 
        if btt_gerar_doc:
            st.session_state["gerar_dec_conteudo"] = True
            
    # EXCLUIR DEC DE CONTEUDO
    if  st.session_state["dados"] and not st.session_state["gerar_dec_conteudo"]:
        btt_excluir_dec = st.button('Excluir Declaração de Conteudo')
        if btt_excluir_dec :
            excluir_dados() 

    # EDITAR ARQUIVO DA DEC E ADICIONA OS PRODUTOS 
    if  st.session_state["gerar_dec_conteudo"]:
        
        # Obtendo a data atual
        data_atual = datetime.now()
        
        # Extraindo o dia, mês e ano
        dia = data_atual.day
        mes = data_atual.month
        ano = data_atual.year

        #DEC_CONTEÚDO_CORREIOS-04_12_2024.docx
        caminho_dec_docx = f"DECS/DEC_CONTEÚDO_CORREIOS-{dia}_{mes}_{ano}.docx"
        
        # Função para editar o documento base
        def criando_arquivo_docx(dados_dos_produtos, caminho_base):
            # Abre o documento base
            doc = Document(caminho_base)
            tabela = doc.tables[0]#pondo que seja a tabela de identificação dos bens
            
            # CALCULA A QUANTIDADE DE LINHAS QUE ESTÃO PREENCHIDAS
            qtd_ln_produtos = 10 + (len(dados_dos_produtos) -1)
            
            # FUNÇÃO QUE EXCLUI AS LINHA QUE NÃO ESTÃO SENDO USADAS
            def excluir_linhas_dec(qtd):
                ultima_linha = 49
                
                while ultima_linha > qtd:
                    # Localiza a tabela específica (no caso, a primeira tabela)
                
                    linha_excluida = tabela.rows[ultima_linha]  # Última linha da tabela
                    
                    # Remover a linha (acessando o XML da tabela)
                    tbl = tabela._tbl  # Acessa o XML da tabela
                    tbl.remove(linha_excluida._tr)  # Remove a linha
                    
                    ultima_linha -= 1          
            excluir_linhas_dec(qtd_ln_produtos)
            
            # FUNÇÃO QUE CENTRALIZA O TEXTO NO CENTRO DA CELULA
            def centralizar_celulas(celula):
                # Centralizar horizontalmente
                for paragrafo in celula.paragraphs:
                    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Centralizar verticalmente
                celula.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

            # FUNÇÃO QUE INSERI OS PRODUTOS NA DEC
            def inserir_produtos(dado, id_linha_indent_bens, item):
                if item < 10:
                    item = f'0{item}'
                else:
                    item=str(item)
                conteudo = dado["conteudo"]
                quantidade = str(dado["quantidade"])
                valor = f'{dado["valor"]:.2f}'
                
                
                linha_tabela2 = tabela.rows[id_linha_indent_bens] 
                for celula in linha_tabela2.cells:
                    #SUBSTITUI O VALOR PADRÃO PELO DADO DO PRODUTO
                    if '##' in celula.text:
                        celula.text = celula.text.replace(f'##',f'{item}')
                        centralizar_celulas(celula)
                    if '@@' in celula.text:
                        celula.text = celula.text.replace('@@',f'{conteudo}')
                    if '**'in celula.text:
                        celula.text = celula.text.replace(f'**',f'{quantidade}')
                        centralizar_celulas(celula)
                    if '&&'in celula.text:
                        celula.text = celula.text.replace('&&', f'{valor}')
                        centralizar_celulas(celula)
            
            # FUNÇÃO QUE INSERI O REMETENTE E O DESTINATARIO NA DEC
            def inserir_dest_rem(id_linha_rem_dest, id_tabela):
                
                # PEGA O DADO DO REMETENTE E DO DESTINATARIO SELECIONADO
                cursor.execute("SELECT * FROM remetente WHERE nome = ?", (remetente_selecionado,))
                dados_remetente = cursor.fetchall()
                
                cursor.execute("SELECT * FROM destinatario WHERE nome = ?", (destinatario_selecionado,))
                dados_destinatario = cursor.fetchall()
                
                tabela = doc.tables[id_tabela]
                linha_tabela1 = tabela.rows[id_linha_rem_dest]
                for celula in linha_tabela1.cells:
                    
                    # SUBSTITUI OS VALORES PADRÕES PELOS DADOS DO BANCO DE DADOS 
                    if '$$' in celula.text:
                        celula.text = celula.text.replace(f'$$',f'{dados_remetente[0][1]}')
                    if '¨¨' in celula.text:
                        celula.text = celula.text.replace(f'¨¨',f'{dados_remetente[0][2]}')
                    if f'%%' in celula.text:
                        celula.text = celula.text.replace(f'%%',f'{dados_remetente[0][3]}')
                    if '==' in celula.text:
                        celula.text = celula.text.replace(f'==',f'{dados_remetente[0][4]}')
                    if '++' in celula.text:
                        celula.text = celula.text.replace(f'++',f'{dados_remetente[0][5]}')
                    if '--' in celula.text:
                        celula.text = celula.text.replace(f'--',f'{dados_remetente[0][6]}')
                    
                    
                    if '//' in celula.text:
                        celula.text = celula.text.replace(f'//',f'{dados_destinatario[0][1]}')
                    if '<<' in celula.text:
                        celula.text = celula.text.replace(f'<<',f'{dados_destinatario[0][2]}')
                    if '>>' in celula.text:
                        celula.text = celula.text.replace(f'>>',f'{dados_destinatario[0][3]}')
                    if '||' in celula.text:
                        celula.text = celula.text.replace(f'||',f'{dados_destinatario[0][4]}')
                    if '??' in celula.text:
                        celula.text = celula.text.replace(f'??',f'{dados_destinatario[0][5]}')
                    if '..' in celula.text:
                        celula.text = celula.text.replace(f'..',f'{dados_destinatario[0][6]}')
               
            # INSEIR A DATA DE ESCRITA DA DEC      
            def inserir_data(id_linha_data):
                data_atual = datetime.now()
        
                # Extraindo o dia, mês e ano
                dia = data_atual.day
                mes = data_atual.month
                ano = data_atual.year
                
                nome_mes = {
                    1: 'Janeiro',
                    2: 'Fevereiro',
                    3: 'Março',
                    4: 'Abril',
                    5: 'Maio',
                    6: 'Junho',
                    7: 'Julho',
                    8: 'Agosto',
                    9: 'Setembro',
                    10: 'Outubro',
                    11: 'Novembro',
                    12: 'Dezembro'
                }
                
                linha_data = tabela.rows[id_linha_data]
                for celula in linha_data.cells:
                    if 'ππ' in celula.text:
                        celula.text = celula.text.replace('ππ',f'{dia}')
                        centralizar_celulas(celula)
                    if 'ΩΩ' in celula.text:
                        celula.text = celula.text.replace('ΩΩ',f'{nome_mes[mes]}')
                        centralizar_celulas(celula)
                    if '∑∑' in celula.text:
                        celula.text = celula.text.replace('∑∑',f'{ano}')
                        centralizar_celulas(celula)
        
            # INSERIR A QTD DE ITENS E O VALOR TOTAL
            def inserir_total(id_linha_total):
                linha_total = tabela.rows[id_linha_total]
                qtd_total = 0
                valor_total = 0
                
                for dado_produto in dados_dos_produtos:
                    qtd_total += dado_produto['quantidade']
                    valor_total += dado_produto['valor']
                    
                for celula in linha_total.cells:
                    # SUBSTITUI OS VALORES PADRÕES PELA QTD DE ITENS E O VALOR TOTAL
                    if 'ββ' in celula.text:
                        celula.text = celula.text.replace('ββ',f'{qtd_total}')
                        centralizar_celulas(celula)
                    if '€€' in celula.text:
                        celula.text = celula.text.replace('€€',f'{valor_total:.2f}')   
                        centralizar_celulas(celula) 
            
            # INSERIR O PESO DA CAIXA
            def inserir_peso(id_linha_peso):
                linha_peso = tabela.rows[id_linha_peso]
                for celula in linha_peso.cells:
                    # INSERE O VALOR PADRÃO PELO PESO DA CAIXA
                    if '®®' in celula.text:
                        celula.text = celula.text.replace('®®',f'{st.session_state["peso"]}kg')
                        centralizar_celulas(celula)
                
                                         
            # Adiciona os dados na tabela
            id_linha_indent_bens = 10
            item = 1
            for dado_produto in dados_dos_produtos:
                inserir_produtos(dado_produto, id_linha_indent_bens, item)
                id_linha_indent_bens+=1
                item +=1
            
            id_linha_total = (len(dados_dos_produtos) + 10)
            for i in range(3):
                inserir_total(id_linha_total)
                id_linha_total += 1
                
            inserir_peso(len(dados_dos_produtos) + 11)
            
            id_linha_rem_dest = 3
            for i in range(5):
                inserir_dest_rem(id_linha_rem_dest, 0)
                id_linha_rem_dest +=1
                i+=1
                
            id_linha_rem_dest = 3
            for i in range(5):
                inserir_dest_rem(id_linha_rem_dest, 1)
                id_linha_rem_dest +=1
                i+=1
                
            id_linha_data = (len(dados_dos_produtos) + 6 + 9)
            inserir_data(id_linha_data)
            
            doc.save(caminho_dec_docx)
            return caminho_dec_docx
        
        caminho_documento_base = "DEC_CONTEUDO_CORREIOS.docx"

        caminho_gerado = criando_arquivo_docx(st.session_state["dados"], caminho_documento_base)
        
        # FUNÇÃO QUE CONVERTE O ARQUIVO DOCX EM PDF
        def docx_para_pdf():
            LIBRE_OFFICE = r'C:\Users\User\Documents\dec_conteudo\LibreOffice\program\soffice.exe'

            arquivo_de_entrada = f'{caminho_dec_docx}'
            pasta_destino = 'DECS'

            p = Popen([LIBRE_OFFICE, '--headless', '--convert-to', 'pdf', '--outdir', pasta_destino, arquivo_de_entrada])
            p.communicate()
        docx_para_pdf()

        # FUNÇÃO QUE FAZ O DOWLOAD DO ARQUIVO 
        def dowload():
            caminho_pdf = caminho_dec_docx.replace('.docx', '.pdf')
            with open(caminho_pdf, "rb") as pdf_file:
                pdf_bytes =pdf_file.read()
                btt_dowload = st.download_button(
                    label="Baixar PDF",
                    data=pdf_bytes,
                    file_name=path.basename(caminho_pdf),
                    mime="application/pdf"
                )    
                
                if btt_dowload:
                    st.session_state["nova_dec"] = True
                    st.rerun()           
        dowload()
        
        st.success("Declaração gerada com sucesso!")
    
    # REMOVER ALGUM ITEM DA DEC
    if st.session_state["dados"] and not st.session_state["gerar_dec_conteudo"]:
        st.subheader('Remover produto da DEC')
        indice = st.number_input('Indice do produto que deseja remover da DEC:', min_value=0)
        btt_remover_produto_indice = st.button(f'Remover produto')
        if btt_remover_produto_indice:
            remover_produto(indice)

    # TABELA DOS DADOS DA DEC
    if st.session_state["dados"] and not st.session_state["gerar_dec_conteudo"]:
        st.subheader('Pré visualização da DEC')
        # Criar DataFrame
        df = pd.DataFrame(st.session_state["dados_tabela"])

        # Exibir tabela interativa
        st.dataframe(df, use_container_width=True)
        
        st.write(f'Peso da Caixa: {st.session_state["peso"]}Kg')
    
# NOVA DEC == TRUE:
else:
    btt_nova_dec = st.button('Nova DEC')
    
    if btt_nova_dec:
        excluir_dados()
        
conn.commit()        
cursor.close()
